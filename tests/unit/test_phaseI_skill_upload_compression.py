import shutil
import sys
import unittest
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from tencent_doc_review.access import DownloadFormat, MCPDownloadPayload, MCPUploadPayload, TencentDocReference, UploadTarget
from tencent_doc_review.skill import SkillRequest
from tencent_doc_review.workflows import SkillPipeline


class _CompressionAwareClient:
    def __init__(self, root: Path) -> None:
        self.root = root
        self.uploaded_local_path: Path | None = None

    async def export_document(self, reference, download_format=DownloadFormat.DOCX):
        image_path = self.root / "large-image.png"
        image = Image.effect_noise((2600, 1800), 120)
        image.save(image_path, format="PNG")

        source_path = self.root / "downloaded.docx"
        document = Document()
        document.add_heading(reference.title, level=1)
        document.add_paragraph("Oversized upload candidate.")
        document.add_picture(str(image_path), width=Inches(6.5))
        document.save(source_path)

        return MCPDownloadPayload(
            reference=reference,
            format=download_format,
            filename=source_path.name,
            source_path=source_path,
            metadata={"source": "compression-aware-client", "used_text_fallback": False, "source_path": str(source_path)},
        )

    async def upload_document(self, local_path, target, remote_filename):
        self.uploaded_local_path = Path(local_path)
        return MCPUploadPayload(
            target=target,
            uploaded_name=remote_filename,
            remote_file_id="remote-compressed-123",
            remote_url="https://docs.qq.com/mock/remote-compressed-123",
            metadata={"source": "compression-aware-client"},
        )


class PhaseISkillUploadCompressionTests(unittest.IsolatedAsyncioTestCase):
    async def test_skill_pipeline_compresses_before_upload_when_threshold_exceeded(self):
        tmpdir = PROJECT_ROOT / "tests" / ".tmp" / f"phasei-skill-{uuid.uuid4().hex}"
        tmpdir.mkdir(parents=True, exist_ok=True)
        try:
            client = _CompressionAwareClient(tmpdir)
            request = SkillRequest(
                source_document=TencentDocReference(doc_id="doc-oversized", title="Oversized Demo"),
                target_location=UploadTarget(folder_id="folder-456", space_type="personal_space"),
                max_upload_size_bytes=300_000,
            )

            response = await SkillPipeline().run(client, request)

            self.assertTrue(response.success)
            self.assertIsNotNone(client.uploaded_local_path)
            assert client.uploaded_local_path is not None
            self.assertTrue(client.uploaded_local_path.name.endswith("-compressed.docx"))
            self.assertTrue(response.metadata["compression_applied"])
            self.assertLessEqual(response.metadata["compression_result_size"], response.metadata["compression_original_size"])
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
