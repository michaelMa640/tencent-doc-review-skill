import shutil
import sys
import unittest
import uuid
import zipfile
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from tencent_doc_review.document import WordAnnotation, WordAnnotator, WordParser


class PhaseCWordAnnotationFlowTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmpdir = PROJECT_ROOT / "tests" / ".tmp" / f"phasec-word-{uuid.uuid4().hex}"
        self.tmpdir.mkdir(parents=True, exist_ok=True)

    def tearDown(self) -> None:
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _create_sample_docx(self) -> Path:
        source_path = self.tmpdir / "source.docx"
        document = Document()
        document.add_heading("示例文档", level=1)
        document.add_paragraph("")
        document.add_paragraph("第一段内容。")
        document.add_paragraph("第二段内容。")
        document.save(source_path)
        return source_path

    def test_word_parser_extracts_paragraphs_and_headings(self):
        source_path = self._create_sample_docx()

        parsed = WordParser().parse(source_path)

        self.assertEqual(parsed.title, "示例文档")
        self.assertEqual(len(parsed.paragraphs), 4)
        self.assertTrue(parsed.paragraphs[0].is_heading)
        self.assertEqual(parsed.paragraphs[2].text, "第一段内容。")

    def test_word_annotator_exports_native_comments_and_summary_block(self):
        source_path = self._create_sample_docx()
        output_path = self.tmpdir / "annotated.docx"

        result = WordAnnotator().annotate(
            source_path=source_path,
            output_path=output_path,
            document_title="示例文档",
            annotations=[
                WordAnnotation(
                    paragraph_index=2,
                    title="事实核查提示",
                    comment="这里需要补充数据来源。",
                    severity="high",
                    source_excerpt="第一段内容。",
                ),
                WordAnnotation(
                    paragraph_index=3,
                    title="总结建议",
                    comment="请在文末补充完整结论。",
                    severity="medium",
                    metadata={"render_mode": "summary_block"},
                ),
            ],
        )

        self.assertTrue(output_path.exists())
        self.assertEqual(result.annotation_count, 2)
        self.assertEqual(result.metadata["comment_mode"], "native+summary")
        self.assertEqual(result.metadata["inline_comment_count"], 1)
        self.assertEqual(result.metadata["summary_block_count"], 1)

        rendered = Document(output_path)
        self.assertEqual(len(rendered.comments), 1)
        self.assertIn("AI Review", [comment.author for comment in rendered.comments])
        self.assertTrue(any("这里需要补充数据来源。" in comment.text for comment in rendered.comments))
        self.assertTrue(any("AI审核总结" == paragraph.text.strip() for paragraph in rendered.paragraphs))
        self.assertTrue(any("总结建议" in paragraph.text for paragraph in rendered.paragraphs))

        with zipfile.ZipFile(output_path) as package:
            names = set(package.namelist())
            self.assertIn("word/comments.xml", names)

    def test_word_annotator_renders_summary_table(self):
        source_path = self._create_sample_docx()
        output_path = self.tmpdir / "annotated-table.docx"

        WordAnnotator().annotate(
            source_path=source_path,
            output_path=output_path,
            document_title="示例文档",
            annotations=[
                WordAnnotation(
                    paragraph_index=3,
                    title="结构模板相关性",
                    comment="模板结构匹配度：75.0%",
                    severity="low",
                    metadata={
                        "render_mode": "summary_block",
                        "summary_table": [
                            ["模板章节", "当前文章是否已有", "当前命中章节名", "状态说明"],
                            ["产品概述", "✓", "产品概述", "已覆盖"],
                            ["竞品对比分析", "✗", "", "缺失"],
                        ],
                    },
                )
            ],
        )

        rendered = Document(output_path)
        self.assertEqual(len(rendered.tables), 1)
        table = rendered.tables[0]
        self.assertEqual(table.cell(0, 0).text, "模板章节")
        self.assertEqual(table.cell(1, 1).text, "✓")
        self.assertEqual(table.cell(2, 3).text, "缺失")

    def test_word_annotator_renders_summary_table_without_preferred_table_style(self):
        source_path = self._create_sample_docx()
        output_path = self.tmpdir / "annotated-table-no-style.docx"

        annotator = WordAnnotator()
        annotator._resolve_summary_table_style = lambda document: None  # type: ignore[method-assign]

        annotator.annotate(
            source_path=source_path,
            output_path=output_path,
            document_title="示例文档",
            annotations=[
                WordAnnotation(
                    paragraph_index=3,
                    title="结构模板相关性",
                    comment="模板结构匹配度：75.0%",
                    severity="low",
                    metadata={
                        "render_mode": "summary_block",
                        "summary_table": [
                            ["模板章节", "当前文章是否已有", "当前命中章节名", "状态说明"],
                            ["产品概述", "✓", "产品概述", "已覆盖"],
                        ],
                    },
                )
            ],
        )

        rendered = Document(output_path)
        self.assertEqual(len(rendered.tables), 1)
        self.assertEqual(rendered.tables[0].cell(1, 1).text, "✓")


if __name__ == "__main__":
    unittest.main()
