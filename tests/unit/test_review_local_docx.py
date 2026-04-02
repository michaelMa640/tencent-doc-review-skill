import json
import shutil
import sys
import unittest
import uuid
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from tencent_doc_review.workflows.skill_pipeline import SkillPipeline


class ReviewLocalDocxTests(unittest.IsolatedAsyncioTestCase):
    async def test_review_local_docx_generates_annotated_docx_report_and_issue_safe_debug_bundle(self):
        tmp_path = PROJECT_ROOT / "downloads" / f"test-review-local-{uuid.uuid4().hex[:8]}"
        tmp_path.mkdir(parents=True, exist_ok=True)
        try:
            source_path = tmp_path / "sample.docx"

            document = Document()
            document.add_heading("产品概述", level=1)
            document.add_paragraph(
                "该产品支持实时动画编辑，并提供多人协作能力。"
                "在调研过程中，我们重点关注其工作流效率、资产管理方式以及团队协同场景下的可用性。"
            )
            document.add_heading("结论与推荐建议", level=1)
            document.add_paragraph(
                "建议继续评估企业级协作场景，并补充与竞品在价格、授权方式和渲染速度上的横向对比。"
                "如果后续要进入团队落地阶段，还需要验证权限管理与资产复用策略。"
            )
            document.save(source_path)

            artifacts = await SkillPipeline().review_local_docx(
                input_path=source_path,
                title="RIVE产品调研报告-Michael",
                provider="mock",
                output_dir=tmp_path,
                debug_output_dir=str(tmp_path / "debug"),
            )

            self.assertTrue(Path(artifacts.annotated_word_path).exists())
            self.assertTrue(Path(artifacts.markdown_report_path).exists())
            self.assertTrue(Path(artifacts.upload_candidate_path).exists())
            self.assertGreaterEqual(artifacts.annotation_count, 1)
            self.assertTrue(artifacts.review_summary)
            self.assertIn("RIVE产品调研报告-Michael-annotated.docx", artifacts.annotated_word_path)

            debug_bundle_path = Path(artifacts.debug_bundle_path)
            self.assertTrue(debug_bundle_path.exists())
            self.assertEqual(debug_bundle_path.suffix, ".json")
            self.assertIn("tdr-debug-", debug_bundle_path.name)
            self.assertIn("RIVE产品调研报告-Michael", debug_bundle_path.name)

            payload = json.loads(debug_bundle_path.read_text(encoding="utf-8"))
            self.assertEqual(payload["bundle_version"], 1)
            self.assertIn("redaction_notice", payload)
            self.assertEqual(payload["config_snapshot"]["llm_provider"], "deepseek")
            self.assertNotIn(str(PROJECT_ROOT), json.dumps(payload, ensure_ascii=False))
        finally:
            shutil.rmtree(tmp_path, ignore_errors=True)

    async def test_review_local_docx_rejects_incomplete_source_document(self):
        tmp_path = PROJECT_ROOT / "downloads" / f"test-review-local-{uuid.uuid4().hex[:8]}"
        tmp_path.mkdir(parents=True, exist_ok=True)
        try:
            source_path = tmp_path / "sample-short.docx"

            document = Document()
            document.add_heading("Reela 产品调研报告", level=1)
            document.add_paragraph("这是一句被错误下载出来的内容。")
            document.save(source_path)

            with self.assertRaisesRegex(ValueError, "Downloaded source document appears incomplete"):
                await SkillPipeline().review_local_docx(
                    input_path=source_path,
                    title="Reela 产品调研报告-Michael",
                    provider="mock",
                    output_dir=tmp_path,
                    debug_output_dir=str(tmp_path / "debug"),
                )
        finally:
            shutil.rmtree(tmp_path, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
