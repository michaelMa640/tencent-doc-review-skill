import json
import shutil
import sys
import uuid
import unittest
from pathlib import Path

from click.testing import CliRunner
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from tencent_doc_review.access import DownloadFormat, MCPDownloadPayload, MCPUploadPayload, TencentDocReference, UploadTarget
from tencent_doc_review.cli import main
from tencent_doc_review.skill import SkillRequest
from tencent_doc_review.workflows import SkillPipeline


class _FakeSkillClient:
    def __init__(self, root: Path) -> None:
        self.root = root

    async def export_document(self, reference, download_format=DownloadFormat.DOCX):
        source_path = self.root / "source.docx"
        document = Document()
        document.add_heading(reference.title, level=1)
        document.add_paragraph("Skill pipeline example content.")
        document.save(source_path)
        return MCPDownloadPayload(
            reference=reference,
            format=download_format,
            filename=source_path.name,
            content_bytes=source_path.read_bytes(),
            source_path=source_path,
            metadata={"source": "fake-skill-client"},
        )

    async def upload_document(self, local_path, target, remote_filename):
        return MCPUploadPayload(
            target=target,
            uploaded_name=remote_filename,
            remote_file_id="remote-123",
            remote_url="https://docs.qq.com/mock/remote-123",
            metadata={"source": "fake-skill-client"},
        )


class PhaseESkillWorkflowTests(unittest.IsolatedAsyncioTestCase):
    async def test_skill_pipeline_returns_unified_skill_response(self):
        tmpdir = PROJECT_ROOT / "tests" / ".tmp" / f"phasee-skill-{uuid.uuid4().hex}"
        tmpdir.mkdir(parents=True, exist_ok=True)
        try:
            request = SkillRequest(
                source_document=TencentDocReference(doc_id="doc-123", title="Skill Demo"),
                target_location=UploadTarget(folder_id="folder-456", path_hint="/审核结果"),
            )

            response = await SkillPipeline().run(_FakeSkillClient(tmpdir), request)

            self.assertTrue(response.success)
            self.assertEqual(response.remote_file_id, "remote-123")
            self.assertTrue(response.annotated_word_path.endswith("-annotated.docx"))
            self.assertEqual(response.target_location.folder_id, "folder-456")
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    def test_cli_skill_info_prints_runtime_json(self):
        runner = CliRunner()
        result = runner.invoke(main, ["skill-info"])

        self.assertEqual(result.exit_code, 0, msg=result.output)
        payload = json.loads(result.output)
        self.assertIn("platform", payload)
        self.assertIn("temp_root", payload)

    def test_cli_skill_run_prints_response(self):
        runner = CliRunner()
        result = runner.invoke(
            main,
            [
                "skill-run",
                "--doc-id",
                "doc-123",
                "--title",
                "Skill Demo",
                "--target-folder-id",
                "folder-456",
                "--target-path",
                "/审核结果",
            ],
        )

        self.assertEqual(result.exit_code, 0, msg=result.output)
        payload = json.loads(result.output)
        self.assertTrue(payload["success"])
        self.assertEqual(payload["remote_file_id"], "mock-remote-file-id")


if __name__ == "__main__":
    unittest.main()
