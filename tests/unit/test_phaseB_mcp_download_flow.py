import shutil
import sys
import unittest
import uuid
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from tencent_doc_review.access import DownloadFormat, DownloadManager, MCPDownloadPayload, TencentDocReference


class _TextFallbackClient:
    async def export_document(self, reference, download_format=DownloadFormat.DOCX):
        suffix = ".docx" if download_format is DownloadFormat.DOCX else ".md"
        return MCPDownloadPayload(
            reference=reference,
            format=download_format,
            filename=f"{reference.title}{suffix}",
            text_content="First paragraph.\n\nSecond paragraph.",
            metadata={"source": "fake-mcp", "used_text_fallback": True},
        )


class _DownloadedFileClient:
    def __init__(self, root: Path) -> None:
        self.root = root

    async def export_document(self, reference, download_format=DownloadFormat.DOCX):
        source_path = self.root / "downloaded-source.docx"
        document = Document()
        document.add_heading(reference.title, level=1)
        document.add_paragraph("Downloaded file paragraph.")
        document.save(source_path)
        return MCPDownloadPayload(
            reference=reference,
            format=download_format,
            filename=source_path.name,
            source_path=source_path,
            metadata={"source": "fake-mcp", "used_text_fallback": False, "source_path": str(source_path)},
        )


class PhaseBMCPDownloadFlowTests(unittest.IsolatedAsyncioTestCase):
    async def test_download_manager_materializes_docx_from_text_fallback(self):
        tmpdir = PROJECT_ROOT / "tests" / ".tmp" / f"phaseb-download-{uuid.uuid4().hex}"
        tmpdir.mkdir(parents=True, exist_ok=True)
        try:
            manager = DownloadManager(root_dir=tmpdir)
            reference = TencentDocReference(doc_id="doc-123", title="Example Review Document")

            downloaded = await manager.download_via_mcp(
                _TextFallbackClient(),
                reference,
                purpose="document",
                download_format=DownloadFormat.DOCX,
            )

            self.assertTrue(downloaded.file_path.exists())
            self.assertEqual(downloaded.file_path.suffix, ".docx")
            paragraphs = [p.text for p in Document(downloaded.file_path).paragraphs]
            self.assertIn("First paragraph.", paragraphs)
            self.assertIn("Second paragraph.", paragraphs)
            self.assertTrue(downloaded.metadata["used_text_fallback"])
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    async def test_download_manager_prefers_existing_source_path(self):
        tmpdir = PROJECT_ROOT / "tests" / ".tmp" / f"phaseb-source-{uuid.uuid4().hex}"
        tmpdir.mkdir(parents=True, exist_ok=True)
        try:
            manager = DownloadManager(root_dir=tmpdir)
            reference = TencentDocReference(doc_id="doc-456", title="Downloaded Original")

            downloaded = await manager.download_via_mcp(
                _DownloadedFileClient(tmpdir),
                reference,
                purpose="document",
                download_format=DownloadFormat.DOCX,
            )

            self.assertTrue(downloaded.file_path.exists())
            self.assertEqual(downloaded.file_path, tmpdir / "downloaded-source.docx")
            paragraphs = [p.text for p in Document(downloaded.file_path).paragraphs]
            self.assertIn("Downloaded file paragraph.", paragraphs)
            self.assertFalse(downloaded.metadata["used_text_fallback"])
            self.assertTrue(downloaded.metadata["source_path"].endswith("downloaded-source.docx"))
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    def test_build_plan_adds_purpose_suffix_for_template_download(self):
        tmpdir = PROJECT_ROOT / "tests" / ".tmp" / f"phaseb-template-{uuid.uuid4().hex}"
        tmpdir.mkdir(parents=True, exist_ok=True)
        try:
            manager = DownloadManager(root_dir=tmpdir)
            reference = TencentDocReference(doc_id="doc-123", title="Quarterly Plan")

            plan = manager.build_plan(
                reference,
                download_format=DownloadFormat.DOCX,
                purpose="template",
            )

            self.assertEqual(plan.filename, "Quarterly-Plan-template.docx")
            self.assertEqual(plan.target_dir.name, "doc-123")
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
