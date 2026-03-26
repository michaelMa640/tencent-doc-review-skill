import shutil
import sys
import unittest
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = PROJECT_ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from tencent_doc_review.document import DocxCompressor


class PhaseHDocxCompressorTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmpdir = PROJECT_ROOT / "tests" / ".tmp" / f"phaseh-compress-{uuid.uuid4().hex}"
        self.tmpdir.mkdir(parents=True, exist_ok=True)

    def tearDown(self) -> None:
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _create_oversized_docx(self) -> Path:
        image_path = self.tmpdir / "large-image.png"
        image = Image.effect_noise((2600, 1800), 120)
        image.save(image_path, format="PNG")

        source_path = self.tmpdir / "source.docx"
        document = Document()
        document.add_heading("Compression Demo", level=1)
        document.add_paragraph("This document contains a large embedded image.")
        document.add_picture(str(image_path), width=Inches(6.5))
        document.save(source_path)
        return source_path

    def test_docx_compressor_reduces_file_size(self):
        source_path = self._create_oversized_docx()
        output_path = self.tmpdir / "compressed.docx"

        result = DocxCompressor().compress_to_target(
            source_path=source_path,
            output_path=output_path,
            target_max_bytes=source_path.stat().st_size - 1024,
        )

        self.assertTrue(output_path.exists())
        self.assertLess(result.compressed_size, result.original_size)
        self.assertTrue(result.changed_entries)


if __name__ == "__main__":
    unittest.main()
