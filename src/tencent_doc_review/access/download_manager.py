"""Download planning and local staging for MCP-derived documents."""

from __future__ import annotations

import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Optional

from docx import Document

from .mcp_adapter import DownloadFormat, MCPDocumentClient, MCPDownloadPayload, TencentDocReference

_SAFE_CHARS_PATTERN = re.compile(r"[^A-Za-z0-9._-]+")


@dataclass
class DownloadPlan:
    """Resolved local download path information."""

    reference: TencentDocReference
    target_dir: Path
    filename: str
    file_path: Path
    download_format: DownloadFormat
    purpose: str = "document"


@dataclass
class DownloadedDocument:
    """Materialized document on local disk."""

    reference: TencentDocReference
    file_path: Path
    download_format: DownloadFormat
    filename: str
    purpose: str = "document"
    metadata: Dict[str, object] = field(default_factory=dict)


class DownloadManager:
    """Create stable local staging paths for MCP-derived documents."""

    def __init__(self, root_dir: Optional[Path] = None) -> None:
        default_root = Path(tempfile.gettempdir()) / "tencent-doc-review" / "downloads"
        self.root_dir = Path(root_dir) if root_dir is not None else default_root

    async def download_via_mcp(
        self,
        client: MCPDocumentClient,
        reference: TencentDocReference,
        purpose: str = "document",
        download_format: DownloadFormat = DownloadFormat.DOCX,
    ) -> DownloadedDocument:
        payload = await client.export_document(reference, download_format=download_format)
        plan = self.build_plan(reference, payload.filename, payload.format, purpose=purpose)
        return self.materialize(payload, plan)

    def build_plan(
        self,
        reference: TencentDocReference,
        filename: Optional[str] = None,
        download_format: DownloadFormat = DownloadFormat.DOCX,
        purpose: str = "document",
    ) -> DownloadPlan:
        target_dir = self.root_dir / self._sanitize_name(reference.doc_id)
        target_dir.mkdir(parents=True, exist_ok=True)

        resolved_filename = filename or self._build_filename(reference, download_format, purpose)
        file_path = target_dir / resolved_filename
        return DownloadPlan(
            reference=reference,
            target_dir=target_dir,
            filename=resolved_filename,
            file_path=file_path,
            download_format=download_format,
            purpose=purpose,
        )

    def materialize(self, payload: MCPDownloadPayload, plan: DownloadPlan) -> DownloadedDocument:
        plan.target_dir.mkdir(parents=True, exist_ok=True)
        if payload.source_path is not None and payload.source_path.exists():
            plan.file_path.write_bytes(payload.source_path.read_bytes())
        elif payload.content_bytes:
            plan.file_path.write_bytes(payload.content_bytes)
        else:
            self._materialize_from_text(payload, plan)
        return DownloadedDocument(
            reference=plan.reference,
            file_path=plan.file_path,
            download_format=plan.download_format,
            filename=plan.filename,
            purpose=plan.purpose,
            metadata=dict(payload.metadata),
        )

    def _build_filename(
        self,
        reference: TencentDocReference,
        download_format: DownloadFormat,
        purpose: str,
    ) -> str:
        stem = self._sanitize_name(reference.title or reference.doc_id)
        suffix = self._suffix_for_format(download_format)
        if purpose and purpose != "document":
            return f"{stem}-{self._sanitize_name(purpose)}{suffix}"
        return f"{stem}{suffix}"

    def _suffix_for_format(self, download_format: DownloadFormat) -> str:
        if download_format is DownloadFormat.DOCX:
            return ".docx"
        if download_format is DownloadFormat.MARKDOWN:
            return ".md"
        return ".txt"

    def _materialize_from_text(self, payload: MCPDownloadPayload, plan: DownloadPlan) -> None:
        text_content = payload.text_content or payload.reference.title or payload.reference.doc_id
        if plan.download_format is DownloadFormat.DOCX:
            document = Document()
            title = payload.reference.title or payload.reference.doc_id
            if title:
                document.add_heading(title, level=1)
            paragraphs = [segment.strip() for segment in text_content.replace("\r\n", "\n").split("\n\n")]
            wrote_paragraph = False
            for paragraph in paragraphs:
                if paragraph:
                    document.add_paragraph(paragraph)
                    wrote_paragraph = True
            if not wrote_paragraph:
                document.add_paragraph("")
            document.save(plan.file_path)
            return

        if plan.download_format is DownloadFormat.MARKDOWN:
            plan.file_path.write_text(text_content, encoding="utf-8")
            return

        plan.file_path.write_text(text_content, encoding="utf-8")

    def _sanitize_name(self, value: str) -> str:
        normalized = _SAFE_CHARS_PATTERN.sub("-", value.strip())
        normalized = normalized.strip("-._")
        return normalized or "document"
