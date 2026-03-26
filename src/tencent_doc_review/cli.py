"""Command-line entrypoints for local text analysis workflows."""

from __future__ import annotations

import asyncio
import json
import platform
import tempfile
from pathlib import Path
from typing import Optional

import click

from .config import get_settings
from .access import TencentDocReference, UploadTarget
from .analyzer.document_analyzer import DocumentAnalyzer
from .llm.factory import create_llm_client
from .skill import SkillRequest, SkillRuntimeInfo
from .tencent_doc_client import TencentDocClient
from .workflows import SkillPipeline
from .writers import DocAppendWriter, ReportGenerator


def _read_text(path: Optional[str]) -> Optional[str]:
    if not path:
        return None
    return Path(path).read_text(encoding="utf-8")


def _create_tencent_doc_client() -> TencentDocClient:
    settings = get_settings()
    return TencentDocClient(
        access_token=settings.tencent_docs_token,
        client_id=settings.tencent_docs_client_id,
        open_id=settings.tencent_docs_open_id,
        base_url=settings.tencent_docs_base_url,
        timeout=settings.request_timeout,
        max_retries=settings.tencent_docs_max_retries,
        retry_delay=settings.tencent_docs_retry_delay,
    )


@click.group()
def main() -> None:
    """Tencent document review CLI."""


@main.command("doctor")
def doctor() -> None:
    """Show local configuration status."""
    settings = get_settings()
    click.echo("Configuration")
    click.echo(f"  LLM_PROVIDER: {settings.llm_provider}")
    click.echo(f"  LLM_API_KEY: {'set' if (settings.llm_api_key or settings.deepseek_api_key) else 'missing'}")
    click.echo(f"  TENCENT_DOCS_TOKEN: {'set' if settings.tencent_docs_token else 'missing'}")
    click.echo(f"  TENCENT_DOCS_CLIENT_ID: {'set' if settings.tencent_docs_client_id else 'missing'}")
    click.echo(f"  TENCENT_DOCS_OPEN_ID: {'set' if settings.tencent_docs_open_id else 'missing'}")


@main.command("debug-doc")
@click.option("--doc-id", "doc_id", required=True, type=str)
def debug_doc(doc_id: str) -> None:
    """Print a safe summary of the Tencent Docs API response structure."""
    asyncio.run(_debug_doc(doc_id))


@main.command("debug-converter")
@click.option("--encoded-id", "encoded_id", required=True, type=str)
def debug_converter(encoded_id: str) -> None:
    """Print a safe summary of the Tencent Docs converter response."""
    asyncio.run(_debug_converter(encoded_id))


@main.command("list-files")
@click.option("--folder-id", "folder_id", type=str)
@click.option("--encoded-id", "encoded_id", type=str)
@click.option("--format", "output_format", type=click.Choice(["table", "json"]), default="table")
def list_files(folder_id: Optional[str], encoded_id: Optional[str], output_format: str) -> None:
    """List Tencent Docs files or resolve an encoded URL id into a real fileId."""
    asyncio.run(_list_files(folder_id, encoded_id, output_format))


@main.command("skill-info")
def skill_info() -> None:
    """Print shared skill runtime information for OpenClaw / Claude Code integration."""
    runtime = SkillRuntimeInfo(
        platform=platform.system().lower(),
        temp_root=str(Path(tempfile.gettempdir()) / "tencent-doc-review"),
    )
    click.echo(json.dumps(runtime.__dict__, ensure_ascii=False, indent=2))


@main.command("analyze")
@click.option("--input-file", "input_file", type=click.Path(exists=True))
@click.option("--doc-id", "doc_id", type=str)
@click.option("--template-file", "template_file", type=click.Path(exists=True))
@click.option("--template-doc-id", "template_doc_id", type=str)
@click.option("--output", "output_path", type=click.Path())
@click.option("--format", "output_format", type=click.Choice(["markdown", "json", "html"]), default="markdown")
@click.option("--writeback-mode", "writeback_mode", type=click.Choice(["none", "append"]), default="none")
@click.option("--provider", "provider", type=str)
@click.option("--api-key", "api_key", type=str)
@click.option("--base-url", "base_url", type=str)
@click.option("--model", "model", type=str)
def analyze(
    input_file: Optional[str],
    doc_id: Optional[str],
    template_file: Optional[str],
    template_doc_id: Optional[str],
    output_path: Optional[str],
    output_format: str,
    writeback_mode: str,
    provider: Optional[str],
    api_key: Optional[str],
    base_url: Optional[str],
    model: Optional[str],
) -> None:
    """Analyze a local file or Tencent Docs document."""
    asyncio.run(
        _analyze(
            input_file,
            doc_id,
            template_file,
            template_doc_id,
            output_path,
            output_format,
            writeback_mode,
            provider,
            api_key,
            base_url,
            model,
        )
    )


@main.command("skill-run")
@click.option("--doc-id", "doc_id", required=True, type=str)
@click.option("--title", "title", required=True, type=str)
@click.option("--target-folder-id", "target_folder_id", required=True, type=str)
@click.option("--target-space-id", "target_space_id", default="", type=str)
@click.option("--target-path", "target_path", default="", type=str)
def skill_run(
    doc_id: str,
    title: str,
    target_folder_id: str,
    target_space_id: str,
    target_path: str,
) -> None:
    """Run the shared skill workflow with a local MCP client implementation."""
    asyncio.run(
        _skill_run(
            doc_id=doc_id,
            title=title,
            target_folder_id=target_folder_id,
            target_space_id=target_space_id,
            target_path=target_path,
        )
    )


async def _analyze(
    input_file: Optional[str],
    doc_id: Optional[str],
    template_file: Optional[str],
    template_doc_id: Optional[str],
    output_path: Optional[str],
    output_format: str,
    writeback_mode: str,
    provider: Optional[str],
    api_key: Optional[str],
    base_url: Optional[str],
    model: Optional[str],
) -> None:
    if bool(input_file) == bool(doc_id):
        raise click.UsageError("Provide exactly one of --input-file or --doc-id.")
    if template_file and template_doc_id:
        raise click.UsageError("Provide at most one of --template-file or --template-doc-id.")
    if writeback_mode != "none" and not doc_id:
        raise click.UsageError("Writeback is only supported with --doc-id.")

    settings = get_settings()
    client = create_llm_client(
        provider=provider or settings.llm_provider,
        api_key=api_key or settings.llm_api_key or settings.deepseek_api_key,
        base_url=base_url or settings.llm_base_url or settings.deepseek_base_url,
        model=model or settings.llm_model or settings.deepseek_model,
        timeout=settings.request_timeout,
    )

    doc_client: Optional[TencentDocClient] = None
    writeback_result = None
    try:
        if doc_id:
            doc_client = _create_tencent_doc_client()
            analyzer = DocumentAnalyzer(llm_client=client, mcp_client=doc_client)
            result = await analyzer.analyze_from_tencent_doc(
                file_id=doc_id,
                template_file_id=template_doc_id,
            )
            if writeback_mode == "append":
                writeback_result = await DocAppendWriter().write(doc_client, doc_id, result)
        else:
            analyzer = DocumentAnalyzer(llm_client=client)
            result = await analyzer.analyze(
                document_text=_read_text(input_file) or "",
                template_text=_read_text(template_file),
                document_title=Path(input_file or "").name,
            )

        rendered = ReportGenerator().render(result, output_format)

        if output_path:
            Path(output_path).write_text(rendered, encoding="utf-8")
            click.echo(f"Saved analysis to {output_path}")
        else:
            click.echo(rendered)
        if writeback_result is not None:
            click.echo(f"Writeback mode: {writeback_result.get('mode', 'append')}")
    finally:
        if doc_client is not None:
            await doc_client.close()
        await client.close()


async def _debug_doc(doc_id: str) -> None:
    doc_client = _create_tencent_doc_client()
    try:
        payload = await doc_client.debug_document_response(doc_id)
        click.echo(json.dumps(payload, ensure_ascii=False, indent=2))
    finally:
        await doc_client.close()


async def _debug_converter(encoded_id: str) -> None:
    doc_client = _create_tencent_doc_client()
    try:
        payload = await doc_client.debug_converter_response(encoded_id)
        click.echo(json.dumps(payload, ensure_ascii=False, indent=2))
    finally:
        await doc_client.close()


async def _list_files(folder_id: Optional[str], encoded_id: Optional[str], output_format: str) -> None:
    if bool(folder_id) == bool(encoded_id):
        raise click.UsageError("Provide exactly one of --folder-id or --encoded-id.")

    doc_client = _create_tencent_doc_client()
    try:
        if encoded_id:
            file_id = await doc_client.convert_encoded_id_to_file_id(encoded_id)
            payload = {
                "encoded_id": encoded_id,
                "file_id": file_id,
                "analyze_command": f'tencent-doc-review analyze --doc-id "{file_id}" --output report.md',
            }
            click.echo(json.dumps(payload, ensure_ascii=False, indent=2))
            return

        items = await doc_client.list_documents(folder_id or "")
        if output_format == "json":
            payload = [
                {
                    "file_id": item.file_id,
                    "title": item.title,
                    "type": item.doc_type,
                    "url": item.url,
                }
                for item in items
            ]
            click.echo(json.dumps(payload, ensure_ascii=False, indent=2))
            return

        if not items:
            click.echo("No files returned.")
            return

        click.echo("file_id\ttype\ttitle")
        for item in items:
            click.echo(f"{item.file_id}\t{item.doc_type or '-'}\t{item.title or '-'}")
    finally:
        await doc_client.close()


async def _skill_run(
    doc_id: str,
    title: str,
    target_folder_id: str,
    target_space_id: str,
    target_path: str,
) -> None:
    class _CliMCPClient:
        async def export_document(self, reference, download_format):
            source_path = Path(tempfile.gettempdir()) / "tencent-doc-review" / f"{reference.doc_id}.docx"
            source_path.parent.mkdir(parents=True, exist_ok=True)
            from docx import Document

            document = Document()
            document.add_heading(reference.title or reference.doc_id, level=1)
            document.add_paragraph("Skill workflow placeholder content.")
            document.save(source_path)

            from .access import MCPDownloadPayload

            return MCPDownloadPayload(
                reference=reference,
                format=download_format,
                filename=source_path.name,
                content_bytes=source_path.read_bytes(),
                source_path=source_path,
                metadata={"source": "cli-skill-placeholder"},
            )

        async def upload_document(self, local_path, target, remote_filename):
            from .access import MCPUploadPayload

            return MCPUploadPayload(
                target=target,
                uploaded_name=remote_filename,
                remote_file_id="mock-remote-file-id",
                remote_url=f"https://docs.qq.com/mock/{remote_filename}",
                metadata={"source": "cli-skill-placeholder"},
            )

    request = SkillRequest(
        source_document=TencentDocReference(doc_id=doc_id, title=title),
        target_location=UploadTarget(
            folder_id=target_folder_id,
            space_id=target_space_id,
            path_hint=target_path,
            display_name=target_path or target_folder_id,
        ),
    )
    response = await SkillPipeline().run(_CliMCPClient(), request)
    click.echo(json.dumps(response.__dict__, ensure_ascii=False, indent=2, default=str))
