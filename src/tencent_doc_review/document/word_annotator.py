"""Write review output back into a `.docx` file."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.run import Run


@dataclass
class WordAnnotation:
    """Annotation to be applied to a Word document."""

    paragraph_index: int
    title: str
    comment: str
    severity: str = "medium"
    source_excerpt: str = ""
    metadata: Dict[str, object] = field(default_factory=dict)


@dataclass
class AnnotatedWordDocument:
    """Result of exporting an annotated Word document."""

    source_path: Path
    output_path: Path
    annotation_count: int
    metadata: Dict[str, object] = field(default_factory=dict)


class WordAnnotator:
    """Apply inline comments and append summary paragraphs at document end."""

    def annotate(
        self,
        source_path: Path,
        output_path: Path,
        annotations: List[WordAnnotation],
        document_title: Optional[str] = None,
    ) -> AnnotatedWordDocument:
        source = Path(source_path)
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        if output.exists():
            output.unlink()

        document = Document(source)
        inline_annotations = [item for item in annotations if item.metadata.get("render_mode") != "summary_block"]
        summary_annotations = [item for item in annotations if item.metadata.get("render_mode") == "summary_block"]

        for annotation in inline_annotations:
            paragraph = self._resolve_paragraph(document, annotation.paragraph_index)
            anchor_run = self._resolve_anchor_run(paragraph)
            document.add_comment(
                runs=anchor_run,
                text=self._build_comment_text(annotation),
                author="AI Review",
                initials="AI",
            )

        if summary_annotations:
            self._append_summary_section(document, summary_annotations)

        document.save(output)
        return AnnotatedWordDocument(
            source_path=source,
            output_path=output,
            annotation_count=len(annotations),
            metadata={
                "document_title": document_title or source.stem,
                "comment_mode": "native+summary" if summary_annotations else "native",
                "inline_comment_count": len(inline_annotations),
                "summary_block_count": len(summary_annotations),
            },
        )

    def _resolve_paragraph(self, document: DocumentObject, paragraph_index: int):
        paragraphs = document.paragraphs
        if not paragraphs:
            return document.add_paragraph("")
        if 0 <= paragraph_index < len(paragraphs):
            return paragraphs[paragraph_index]
        return paragraphs[0]

    def _resolve_anchor_run(self, paragraph) -> Run:
        if paragraph.runs:
            return paragraph.runs[0]
        return paragraph.add_run(" ")

    def _build_comment_text(self, annotation: WordAnnotation) -> str:
        parts = [f"[{annotation.severity.upper()}] {annotation.title}"]
        body = (annotation.comment or "").strip()
        if body:
            parts.append(body)
        excerpt = (annotation.source_excerpt or "").strip()
        if excerpt:
            parts.append(f"命中原文：{excerpt}")
        return "\n\n".join(parts)

    def _append_summary_section(self, document: DocumentObject, annotations: List[WordAnnotation]) -> None:
        document.add_page_break()
        document.add_heading("AI审核总结", level=1)
        document.add_paragraph("以下内容为整篇层面的审核结论，未挂在具体句子评论中。")

        for index, annotation in enumerate(annotations, start=1):
            title_paragraph = document.add_paragraph()
            title_paragraph.add_run(f"{index}. 【{annotation.title}】").bold = True
            if not annotation.comment:
                continue
            for line in str(annotation.comment).splitlines():
                stripped = line.strip()
                if stripped:
                    document.add_paragraph(stripped)
