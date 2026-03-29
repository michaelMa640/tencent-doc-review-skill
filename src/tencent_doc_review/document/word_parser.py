"""Parse `.docx` files into normalized paragraph and sentence structures."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Dict, List

from docx import Document


@dataclass
class ParagraphNode:
    """Normalized paragraph block extracted from a Word document."""

    index: int
    text: str
    style_name: str = ""
    is_heading: bool = False
    heading_level: int = 0
    metadata: Dict[str, object] = field(default_factory=dict)


@dataclass
class SentenceNode:
    """Original sentence block extracted from a Word paragraph."""

    index: int
    paragraph_index: int
    text: str
    paragraph_text: str = ""
    is_heading: bool = False
    metadata: Dict[str, object] = field(default_factory=dict)


@dataclass
class ParsedWordDocument:
    """Structured representation of a parsed Word document."""

    source_path: Path
    paragraphs: List[ParagraphNode]
    sentences: List[SentenceNode]
    title: str = ""


class WordParser:
    """Read `.docx` files and extract paragraph-level structure."""

    def parse(self, file_path: Path) -> ParsedWordDocument:
        path = Path(file_path)
        document = Document(path)
        paragraphs: List[ParagraphNode] = []
        sentences: List[SentenceNode] = []
        sentence_index = 0

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style is not None else ""
            heading_level = self._detect_heading_level(style_name, text)
            is_heading = heading_level > 0
            paragraphs.append(
                ParagraphNode(
                    index=index,
                    text=text,
                    style_name=style_name,
                    is_heading=is_heading,
                    heading_level=heading_level,
                )
            )
            for sentence in self._split_sentences(text, is_heading=is_heading):
                sentences.append(
                    SentenceNode(
                        index=sentence_index,
                        paragraph_index=index,
                        text=sentence,
                        paragraph_text=text,
                        is_heading=is_heading,
                    )
                )
                sentence_index += 1

        title = next((node.text for node in paragraphs if node.text), path.stem)
        return ParsedWordDocument(
            source_path=path,
            paragraphs=paragraphs,
            sentences=sentences,
            title=title,
        )

    def _detect_heading_level(self, style_name: str, text: str) -> int:
        normalized_style = style_name.lower().strip()
        if normalized_style == "title":
            return 1

        if normalized_style.startswith("heading"):
            parts = normalized_style.split()
            if parts and parts[-1].isdigit():
                return max(1, int(parts[-1]))
            return 1

        stripped = text.strip()
        if not stripped:
            return 0

        if len(stripped) <= 30 and (
            stripped.startswith(tuple(f"{i}." for i in range(1, 10)))
            or stripped.startswith(tuple(f"{i}、" for i in range(1, 10)))
        ):
            return 2

        return 0

    def _split_sentences(self, text: str, is_heading: bool = False) -> List[str]:
        stripped = text.strip()
        if not stripped:
            return []
        if is_heading:
            return [stripped]

        parts = re.split(r"(?<=[。！？!?；;])\s+|(?<=[。！？!?；;])|(?<=\.)\s+(?=[A-Z])|\n+", stripped)
        sentences = [part.strip() for part in parts if part.strip()]
        return sentences or [stripped]
